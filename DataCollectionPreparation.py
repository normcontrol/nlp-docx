import warnings
import re
from docx import Document
import pandas as pd
from nltk.corpus import stopwords
import pymorphy2
morph = pymorphy2.MorphAnalyzer(lang='ru')
import string
import logging

class DataCollectionPreparation(object):
    def __init__(self, path):
        logging.basicConfig(level=logging.DEBUG, filename='log.log', format='%(asctime)s %(levelname)s:%(message)s')
        logging.debug("DataCollectionPreparation run")
        self.russian_stopwords = stopwords.words("russian")
        self. russian_stopwords = set(self.russian_stopwords + ['который', 'таблица', 'рисунок',
                                                     'тот', 'также', 'этот', 'это',
                                                     'такой', 'каждый', 'другой'])
        self.path = path

    def get_data(self):
        '''

        :return:
        '''
        logging.debug("DataCollectionPreparation.get_data run")
        lst_dict_sections_texts = []
        lst_papers = []
        warnings.filterwarnings('ignore')
        f = open(self.path, 'rb')
        document = Document(f) # чтение pdf модулем docx
        paper = [self.__simple_sentence_preproc(p.text) for p in document.paragraphs]
        paper = [sent for sent in paper if sent != '']
        lst_papers.append(paper)
        lst_dict_sections_texts.append(self.__get_docx_sections_texts(document))
        df_docs = pd.DataFrame()
        # df_docs['path_doc'] = lst_paths_docs
        # df_docs['title'] = df_docs['path_doc'].apply(lambda x: x.split('/')[-1])
        df_docs['text_paper'] = lst_papers

        df_docs['dict_sections_texts'] = lst_dict_sections_texts

        # df_docs['faculty_department'] = df_docs['path_doc'].apply(lambda x: x.replace(path_docs, '').split('/')[:-1])
        # df_docs['faculty'] = df_docs['faculty_department'].apply(lambda x: x[2])  # apply(lambda x: x[0])
        # df_docs['department'] = df_docs['faculty_department'].apply(lambda x: x[1] if len(x) > 1 else x[0])

        df_docs['degree'] = ['bachelor' for _ in range(len(df_docs))]
        df_docs['main_preproc_text'] = df_docs['dict_sections_texts'].map(self.__get_main_preproc_text)
        logging.debug("DataCollectionPreparation.get_data df_docs is ready.")
        # return {'hello': path}
        return df_docs

    def __simple_sentence_preproc(self, sentence):
        '''
        :param self:
        :param sentence: Обработка пораграфов документов
        :return: Обработанные параграфы
        '''
        sentence = sentence.replace('\t', ' ').replace('\n', ' ').replace(' ..', '.')
        sentence = self.__all_to_first_upper_letter_of_words(sentence)
        return ''.join(re.findall('[А-Яа-я\. ]', sentence)).replace('  ', ' ').replace(' .', '.').strip()

    def __all_to_first_upper_letter_of_words(self, sent):
        '''
        Замена \n и \t на пробелы
        :param sent: параграф
        :return:
        '''
        sent = sent.replace('\n', ' ').replace('\t', ' ')
        return ' '.join([self.__all_to_first_upper_letter(word) for word in sent.split() if word != ''])

    def __all_to_first_upper_letter(self, word):
        '''
        :param word: Каждое слово в приложении
        :return:
        '''
        set_ul_word = list(set([l.isupper() for l in word.replace('.', '')]))
        if len(set_ul_word) == 1 and set_ul_word[0] == True:
            return word[0] + word[1:].lower()
        else:
            return word

    def __get_docx_sections_texts(self, document):
        '''

        :param document:
        :return:
        '''
        dict_sections_texts = {}
        paragraphs = document.paragraphs
        section_name = paragraphs[0].text
        section_name = self.__simple_sentence_preproc(section_name)
        dict_sections_texts[section_name] = []
        n = 1
        for line, p in enumerate(paragraphs[1:]):
            section_alignment = str(p.paragraph_format.alignment).split()[0]
            section_text = p.text
            section_text = self.__simple_sentence_preproc(section_text)

            # print(n, section_alignment, section_text)
            if section_alignment == 'CENTER':
                section_name = section_text
                dict_sections_texts[section_name] = []

            elif section_alignment != 'CENTER':
                if section_text != '':
                    dict_sections_texts[section_name].append(section_text)

            n = n + 1

        dict_sections_texts_ = {}
        for key, lst_sents in dict_sections_texts.items():
            key = key.lower()
            if key == '' or 'картинка' in key or 'рисунок' in key or 'таблица' in key:
                continue
            dict_sections_texts_[key] = ' '.join(lst_sents).replace(' ..', '.').replace(' .', '.')

        return dict_sections_texts_

    def __sentence_preproc(self, sentence):
        '''

        :param sentence:
        :return:
        '''
        sentence = ''.join([ch for ch in sentence if ch not in string.punctuation])
        sentence = re.sub(r'[^а-яА-Я]', ' ', sentence).strip().replace('  ', '')
        sentence = sentence.split()

        sentence_new = [morph.parse(word)[0].normal_form for word in sentence if word not in self.russian_stopwords and
                        str(morph.parse(word)[0].tag) != 'UNKN']
        sentence_new = ' '.join(sentence_new)
        return sentence_new

    def __get_main_preproc_text(self, dict_sections_texts):
        '''

        :param dict_sections_texts:
        :return:
        '''
        text = ''
        for sents in dict_sections_texts.values():
            if sents != '':
                text = text + " " + sents

        return self.__sentence_preproc(text)